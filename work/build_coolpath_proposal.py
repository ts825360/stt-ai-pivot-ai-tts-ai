from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = r"C:\Users\ts825\Documents\Codex\2026-07-02\stt-ai-pivot-ai-tts-ai\outputs\CoolPath_AI_기획서.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 96, 108)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "D8DEE8"


def set_run_font(run, name="Calibri", east_asia="Malgun Gothic", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", east_asia="Malgun Gothic", size=None, color=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_first_row_as_header(table):
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def paragraph(text="", style=None, align=None, before=None, after=None, line_spacing=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    return p


def add_body(text, after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    run = p.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))
    return p


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_callout(title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    set_table_borders(table, color="D8DEE8")
    mark_first_row_as_header(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11.5, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=RGBColor(30, 30, 30))
    paragraph(after=8)
    return table


def fill_cell(cell, text, bold=False, color=None, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, color=color or RGBColor(0, 0, 0), bold=bold)


def add_table(headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_dxa)
    set_table_borders(table)
    mark_first_row_as_header(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], TABLE_FILL)
        fill_cell(hdr[idx], header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            fill_cell(cells[idx], item, size=10)
    set_table_width(table, widths_dxa)
    paragraph(after=8)
    return table


doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
set_style_font(styles["Normal"], size=11, color=RGBColor(0, 0, 0))
styles["Normal"].paragraph_format.space_after = Pt(8)
styles["Normal"].paragraph_format.line_spacing = 1.333
set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
for style_name in ("List Bullet", "List Number"):
    set_style_font(styles[style_name], size=11, color=RGBColor(0, 0, 0))
    styles[style_name].paragraph_format.left_indent = Inches(0.375)
    styles[style_name].paragraph_format.first_line_indent = Inches(-0.194)
    styles[style_name].paragraph_format.space_after = Pt(4)
    styles[style_name].paragraph_format.line_spacing = 1.208

header = section.header.paragraphs[0]
header.text = ""
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header.add_run("CoolPath AI | Project Proposal")
set_run_font(hr, size=9, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("CoolPath AI 기획서")
set_run_font(fr, size=9, color=MUTED)

# Cover page: proposal_centerpiece pattern
paragraph("AI 기반 파리 여행 이동 전략 추천 서비스", align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("CoolPath AI")
set_run_font(r, size=28, color=INK, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("오늘 나에게 가장 적합한 이동 전략을 추천하는 실감형 AI 이동 보조 앱")
set_run_font(r, size=14, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(26)
r = p.add_run("기획서")
set_run_font(r, size=11, color=MUTED, bold=True)

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
set_table_width(meta, [2100, 7260])
set_table_borders(meta)
mark_first_row_as_header(meta)
metadata_rows = [
    ("프로젝트명", "CoolPath AI"),
    ("서비스 정의", "AI 기반 파리 여행 이동 전략 추천 서비스"),
    ("대상 사용자", "파리 도심을 여행하는 외국인 관광객 및 장시간 도보 이동이 부담되는 사용자"),
    ("핵심 차별점", "최단 경로가 아니라 현재 날씨, 시간대, 도보 피로도, 주변 공간 조건을 고려한 이동 전략 추천"),
]
for i, (label, value) in enumerate(metadata_rows):
    set_cell_shading(meta.cell(i, 0), TABLE_FILL)
    fill_cell(meta.cell(i, 0), label, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(meta.cell(i, 1), value)
set_table_width(meta, [2100, 7260])

paragraph(after=18)
add_callout(
    "기획 핵심 문장",
    "기존 네비게이션은 '가장 빠른 길'을 찾습니다. CoolPath AI는 '오늘 나에게 가장 적합한 이동 전략'을 추천합니다.",
)

doc.add_page_break()

add_heading("1. 프로젝트 개요", 1)
add_body(
    "CoolPath AI는 파리 주요 추천 여행지 사이를 이동하는 사용자를 위해 현재 날씨, 시간대, 거리, 도로 환경, 강가 여부, 대중교통 활용 가능성을 종합하여 가장 쾌적한 이동 전략을 추천하는 AI 기반 서비스이다. 사용자는 출발지와 도착지를 파리 내 주요 관광지 목록에서 선택하고, 앱은 후보 경로의 이동시간과 도보거리뿐 아니라 체감 더위와 야외 노출 가능성을 함께 평가한다."
)
add_body(
    "본 프로젝트는 기존 지도 서비스처럼 단순히 최단 경로를 제공하는 것을 목표로 하지 않는다. 사용자가 실제 여행 현장에서 느끼는 더위, 피로, 햇빛 노출, 이동 부담을 줄이는 방향으로 이동 방식을 판단하고, 왜 해당 경로가 적합한지 AI가 설명하는 것을 핵심 가치로 한다."
)

add_heading("2. 기획 배경 및 문제 정의", 1)
add_body(
    "파리 여행자는 에펠탑, 루브르 박물관, 오르세 미술관, 개선문, 몽마르트르 등 여러 관광지를 하루 안에 이동하는 경우가 많다. 지도 앱은 이동시간과 거리 중심으로 경로를 제공하지만, 폭염이나 높은 습도, 강한 자외선, 낮은 풍속처럼 실제 이동 피로도에 영향을 주는 조건은 충분히 반영하지 못한다."
)
add_bullet("같은 30분 이동이라도 도보 3km와 지하철 이동 후 도보 500m는 체감 피로도가 다르다.")
add_bullet("35°C 이상의 고온, 높은 습도, 약한 바람은 짧은 도보 이동도 훨씬 부담스럽게 만든다.")
add_bullet("강가, 건물 밀집도, 도로폭, 시간대에 따라 실제 햇빛 노출과 체감 쾌적도가 달라질 수 있다.")
add_bullet("여행자는 빠른 길뿐 아니라 오늘의 컨디션과 날씨에 맞는 이동 방식을 필요로 한다.")

add_heading("3. 서비스 목표", 1)
add_body(
    "CoolPath AI의 목표는 사용자가 파리 여행 중 더 안전하고 쾌적하게 이동할 수 있도록 경로별 이동 전략을 비교하고 추천하는 것이다. 특히 폭염, 높은 습도, 긴 도보거리처럼 이동 피로도를 높이는 상황에서 대중교통을 적극적으로 활용하거나, 강가와 건물 밀집 구간을 고려한 상대적으로 쾌적한 경로를 제안한다."
)
add_bullet("파리 주요 관광지 간 이동을 대상으로 구현 범위를 제한하여 실현 가능성을 높인다.")
add_bullet("지도, 날씨, 시간대, 공간 조건을 활용해 경로별 Heat Exposure Score와 Comfort Score를 계산한다.")
add_bullet("AI는 점수 결과를 해석하여 사용자가 이해하기 쉬운 추천 이유를 자연어로 제공한다.")
add_bullet("최단 경로 대비 최대 20분 이내의 우회만 허용하여 쾌적성과 현실성을 함께 확보한다.")

add_heading("4. 서비스 범위", 1)
add_body(
    "본 프로젝트는 모든 도시와 모든 목적지를 지원하는 범용 길찾기 앱이 아니라, 파리 내 추천 여행지를 중심으로 하는 제한형 이동 전략 추천 서비스로 설계한다. 이를 통해 지오코딩, 대중교통 실시간 데이터, 무한한 경로 조합 문제를 줄이고, AI 판단 구조와 사용자 경험을 명확하게 구현할 수 있다."
)
add_table(
    ["구분", "내용"],
    [
        ("공간 범위", "프랑스 파리 시내 및 주요 관광지 주변"),
        ("출발/도착 입력", "자유 입력 대신 추천 여행지 목록 선택 방식"),
        ("추천 여행지 예시", "에펠탑, 루브르 박물관, 오르세 미술관, 개선문, 몽마르트르, 노트르담 대성당, 생트샤펠, 마레 지구, 뤽상부르 공원, 샹젤리제"),
        ("경로 데이터", "후보 경로의 거리, 도보거리, 예상 이동시간, 대중교통 활용 여부를 기반 데이터로 사용"),
        ("실시간 제외 요소", "지하철 도착시간, 실시간 혼잡도, 실시간 버스 배차 정보는 MVP 범위에서 제외"),
    ],
    [2100, 7260],
)

add_heading("5. 핵심 기능", 1)
add_table(
    ["기능", "설명"],
    [
        ("여행지 선택", "파리 내 주요 관광지 목록에서 출발지와 도착지를 선택한다."),
        ("후보 경로 비교", "빠른 경로, 도보 적은 경로, 쾌적 우선 경로 등 여러 후보를 비교한다."),
        ("날씨 기반 판단", "현재 기온, 습도, 풍속을 바탕으로 체감 더위와 이동 피로도를 반영한다."),
        ("태양 위치 추정", "시간대에 따른 해의 방향과 고도를 고려해 직사광선 노출 가능성을 추정한다."),
        ("공간 조건 보정", "도로폭, 주변 건물 밀도, 세느강 인접 여부를 기반으로 구간별 쾌적도를 보정한다."),
        ("Comfort Score", "경로별 이동 쾌적도 점수를 계산하고 가장 적합한 이동 전략을 추천한다."),
        ("AI 설명 생성", "추천 경로가 선택된 이유와 다른 경로가 제외된 이유를 자연어로 설명한다."),
    ],
    [2300, 7060],
)

add_heading("6. AI 판단 구조", 1)
add_body(
    "AI는 직접 지도를 새로 만드는 역할이 아니라, 지도와 날씨 데이터에서 도출된 후보 경로 정보를 해석하고 사용자에게 적합한 이동 전략을 선택하는 역할을 담당한다. 구현은 규칙 기반 점수 계산과 AI 자연어 설명 생성을 결합하는 방식으로 설계한다."
)
add_numbered("사용자가 출발지와 도착지를 선택한다.")
add_numbered("시스템이 후보 경로를 3~4개 생성하거나 미리 준비된 경로 데이터를 불러온다.")
add_numbered("각 경로를 일정 거리 단위의 구간으로 나누고 방향, 거리, 주변 조건을 계산한다.")
add_numbered("현재 시간대의 태양 위치와 날씨 정보를 바탕으로 구간별 열 노출 위험을 추정한다.")
add_numbered("도보거리, 예상 이동시간, 체감온도, 풍속, 도로폭, 건물 밀도, 강가 여부를 종합하여 Comfort Score를 산출한다.")
add_numbered("AI가 가장 점수가 높은 경로를 추천하고, 추천 이유를 사용자가 이해하기 쉬운 문장으로 생성한다.")

add_heading("7. 평가 요소 및 점수화", 1)
add_body(
    "CoolPath AI는 경로를 단순히 빠른 순서로 정렬하지 않고, 사용자가 실제 이동 중 느끼는 부담을 정량화한다. 점수는 정밀한 기상 예보 모델이 아니라, 여행 상황에서 합리적인 이동 판단을 돕는 쾌적도 추정 모델로 사용된다."
)
add_table(
    ["평가 요소", "판단 기준", "점수 반영"],
    [
        ("총 이동시간", "목적지까지 걸리는 전체 시간", "시간이 길수록 패널티"),
        ("도보거리", "야외에 노출되는 총 보행 거리", "길수록 피로도 패널티"),
        ("기온/습도", "고온다습 조건 여부", "체감 더위 상승 시 패널티"),
        ("풍속", "바람에 의한 냉각 가능성", "풍속이 낮으면 패널티, 강가 근처에서는 보너스"),
        ("태양 위치", "시간대별 해의 방향과 고도", "직사광선 노출 가능성이 높으면 패널티"),
        ("도로폭", "넓은 도로와 좁은 골목의 차이", "넓은 도로는 그늘 가능성이 낮아 패널티"),
        ("주변 건물", "건물 밀집 여부", "건물이 많으면 그늘 가능성 보너스"),
        ("강가 여부", "세느강 인접 구간 여부", "바람과 개방감에 따른 쾌적도 보너스"),
        ("대중교통 활용", "도보거리 감소 효과", "더위가 심할 때 보너스"),
        ("우회 제한", "최단 경로 대비 추가 시간", "최대 +20분 이내만 추천 후보로 허용"),
    ],
    [1700, 3960, 3700],
)

add_heading("8. Comfort Score 계산 개념", 1)
add_callout(
    "Comfort Score 개념식",
    "Comfort Score = 100 - 도보거리 패널티 - 체감더위 패널티 - 직사광선 노출 패널티 - 우회시간 패널티 - 환승 부담 + 대중교통 활용 보너스 + 강가/건물 그늘 보정",
)
add_body(
    "예를 들어 현재 조건이 기온 35°C, 습도 72%, 풍속 1m/s, 목적지까지 거리 4.2km라면 CoolPath AI는 장시간 도보 이동의 열 노출 위험을 높게 평가한다. 이 경우 최단 경로가 도보 42분이라도, 지하철이나 버스를 이용해 도보거리를 1km 내외로 줄이는 경로가 더 높은 점수를 받을 수 있다."
)
add_table(
    ["경로", "이동 방식", "예상 시간", "도보거리", "판단 결과"],
    [
        ("A", "전체 도보", "42분", "4.2km", "고온다습 조건에서 장시간 야외 노출이 커서 비추천"),
        ("B", "도보 + 지하철 + 도보", "31분", "1.1km", "도보거리와 전체 시간이 모두 줄어 최우선 추천"),
        ("C", "도보 + 버스 + 도보", "34분", "1.2km", "추천 가능하나 대기/정류장 접근 부담을 고려"),
        ("D", "강가 우회 도보", "58분", "5.0km", "쾌적 구간이 있으나 우회와 도보 부담이 커 보조 후보"),
    ],
    [1000, 2600, 1300, 1300, 3160],
)

add_heading("9. 사용자 모드", 1)
add_body(
    "사용자는 이동 목적에 따라 추천 기준을 선택할 수 있다. 단, 어떤 모드에서도 최단 경로보다 20분 이상 오래 걸리는 경로는 기본 추천에서 제외하여 여행 일정의 현실성을 유지한다."
)
add_table(
    ["모드", "우선 기준", "추천 방식"],
    [
        ("시간 우선", "전체 이동시간, 환승 부담 최소화", "가장 빠른 경로를 중심으로 추천하되 과도한 도보는 경고"),
        ("쾌적 우선", "도보거리 감소, 햇빛 노출 감소, 강가/건물 보정", "최대 +20분 이내에서 더 편한 경로 추천"),
        ("균형 추천", "시간과 쾌적도 균형", "일반 여행자에게 가장 무난한 경로 자동 선택"),
    ],
    [1700, 3460, 4200],
)

add_heading("10. 화면 구성", 1)
add_bullet("메인 화면: 출발 여행지, 도착 여행지, 사용자 모드 선택")
add_bullet("날씨 카드: 현재 기온, 습도, 풍속, 체감 더위, 시간대 표시")
add_bullet("추천 경로 카드: 추천 점수, 예상 시간, 도보거리, 대중교통 활용 여부 표시")
add_bullet("경로 비교 화면: 후보 경로별 Comfort Score와 장단점 비교")
add_bullet("AI 설명 영역: 왜 이 경로를 추천하는지 자연어로 설명")
add_bullet("지도형 화면: 실제 지도 API 또는 지도 스타일 UI 위에 후보 경로 표시")

add_heading("11. 실감미디어 요소", 1)
add_body(
    "본 서비스의 실감미디어 요소는 사용자가 실제 파리 도심을 이동하는 상황을 데이터 기반 시각 정보와 AI 해석으로 확장한다는 점에 있다. 사용자는 단순히 지도 위의 선을 따라가는 것이 아니라, 현재 시간의 해 위치, 체감 더위, 도로 주변 환경, 강가 여부가 반영된 이동 전략을 시각적으로 확인한다."
)
add_body(
    "경로는 이동시간만 표시되는 것이 아니라 열 노출 위험, 쾌적도 점수, 추천 사유가 함께 제공된다. 이를 통해 사용자는 화면 속 경로 정보를 실제 현장의 날씨와 신체 피로도까지 고려한 경험으로 받아들이게 된다."
)

add_heading("12. AI 활용 방식", 1)
add_bullet("경로별 조건 비교 및 점수 해석")
add_bullet("현재 날씨와 이동거리 기반 피로도 판단")
add_bullet("사용자 모드에 따른 가중치 조정")
add_bullet("추천 경로 선택 및 제외 경로 설명")
add_bullet("사용자에게 이해하기 쉬운 자연어 추천 문장 생성")
add_bullet("향후 사용자 피드백을 반영한 개인화 추천 확장")

add_heading("13. 구현 난이도 및 MVP 전략", 1)
add_body(
    "본 프로젝트는 범용 지도 앱으로 구현하면 난이도가 높지만, 파리 주요 추천 여행지로 범위를 제한하면 과제용 MVP로 충분히 구현 가능하다. 핵심은 실시간 교통 데이터를 완벽히 반영하는 것이 아니라, 후보 경로를 비교하고 날씨 및 공간 조건을 기반으로 AI가 이동 전략을 판단하는 흐름을 설득력 있게 보여주는 것이다."
)
add_table(
    ["구현 단계", "내용", "난이도"],
    [
        ("1단계", "파리 추천 여행지 목록과 후보 경로 데이터 구성", "중간"),
        ("2단계", "날씨 데이터 입력 또는 API 연동", "중간"),
        ("3단계", "Comfort Score 계산 로직 구현", "중간"),
        ("4단계", "AI 추천 사유 생성", "중간"),
        ("5단계", "지도형 UI와 경로 카드 구현", "중간~상"),
        ("6단계", "실제 지도 API 및 경로 데이터 확장", "상"),
    ],
    [1600, 5460, 2300],
)

add_heading("14. 한계점 및 보완 방향", 1)
add_body(
    "MVP 단계에서는 실시간 지하철 도착시간, 실시간 혼잡도, 정밀한 건물 그림자 시뮬레이션을 제외한다. 또한 파리 전체의 미세 기후를 거리 단위로 정확히 측정하기보다는, 도시 단위 날씨 데이터와 경로 주변 조건을 조합해 쾌적도를 추정한다."
)
add_bullet("정확한 그림자 계산을 위해서는 건물 높이, 도로폭, 태양 위치, 시간대별 3D 도시 데이터가 필요하다.")
add_bullet("파리 내부에서도 실제 기온과 풍속은 위치마다 달라질 수 있으나, MVP에서는 기본 날씨와 보정값을 활용한다.")
add_bullet("대중교통 실시간 정보가 없으면 배차 지연이나 혼잡도를 반영하기 어렵다.")
add_bullet("향후에는 실시간 교통 API, 사용자 피드백, 웨어러블 기반 피로도 데이터, 정밀 그림자 데이터를 추가할 수 있다.")

add_heading("15. 기대 효과", 1)
add_body(
    "CoolPath AI는 여행자가 더운 날씨에도 무리한 도보 이동을 줄이고, 자신에게 맞는 이동 방식을 선택하도록 돕는다. 특히 낯선 도시에서 지도를 읽는 부담을 줄이고, 현재 상황에 맞는 판단 근거를 제공함으로써 여행 경험의 안전성과 만족도를 높일 수 있다."
)
add_bullet("여행자의 야외 열 노출과 이동 피로도를 줄인다.")
add_bullet("단순 최단 경로 중심의 길찾기에서 벗어나 상황 맞춤형 이동 전략을 제공한다.")
add_bullet("AI 판단 과정을 설명함으로써 사용자가 추천 결과를 신뢰할 수 있게 한다.")
add_bullet("파리 여행이라는 명확한 시나리오를 통해 과제 구현 범위와 발표 설득력을 동시에 확보한다.")

add_heading("16. 결론", 1)
add_body(
    "CoolPath AI는 기존 지도 서비스의 빠른 길 찾기 기능을 대체하려는 서비스가 아니라, 여행자가 실제 현장에서 느끼는 더위와 피로를 고려해 더 적합한 이동 전략을 추천하는 AI 기반 실감형 이동 보조 서비스이다. 파리 주요 여행지라는 제한된 범위를 설정함으로써 구현 가능성을 높이고, 날씨, 시간대, 공간 조건, 대중교통 활용 여부를 결합한 Comfort Score를 통해 차별화된 AI 판단 구조를 제시한다."
)
add_body(
    "따라서 본 프로젝트는 '그늘길 추천 앱'을 넘어, 오늘의 환경과 사용자의 이동 부담을 함께 고려하는 AI 이동 전략 추천 서비스로 발전할 수 있다."
)

doc.save(OUT_PATH)
print(OUT_PATH)
